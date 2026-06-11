import csv
import io
import os
import webbrowser
from datetime import date, datetime, timedelta
from threading import Timer

from docx import Document
from docx.oxml.ns import qn
from flask import Flask, redirect, render_template, request, Response, url_for

import database as db

app = Flask(__name__)

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "documents", "timesheet_template.docx")
_DAY_NAMES_NL = ["Maandag", "Dinsdag", "Woensdag", "Donderdag", "Vrijdag", "Zaterdag", "Zondag"]


def _cell_set(cell, text):
    """Write text into a table cell, clearing any existing runs."""
    para = cell.paragraphs[0]
    for child in list(para._p):
        if child.tag in (qn("w:r"), qn("w:del"), qn("w:ins")):
            para._p.remove(child)
    para.add_run(text)


def _fill_timesheet(intervention):
    doc = Document(TEMPLATE_PATH)
    dt = datetime.strptime(intervention["date"], "%Y-%m-%d")
    week_num = dt.isocalendar()[1]
    weekday = _DAY_NAMES_NL[dt.weekday()]

    hours = intervention["duration_hours"]
    hours_str = str(int(hours)) if hours == int(hours) else str(hours).replace(".", ",")

    # Table 1 – client name
    _cell_set(doc.tables[1].rows[1].cells[1], intervention["client_name"])

    # Table 2 – week number, date, weekday, type as proj.ref
    t2 = doc.tables[2].rows[0]
    _cell_set(t2.cells[1], str(week_num))
    _cell_set(t2.cells[3], dt.strftime("%d/%m/%Y"))
    _cell_set(t2.cells[5], weekday)
    _cell_set(t2.cells[7], intervention["type"])

    # Table 3 – normale uren total
    _cell_set(doc.tables[3].rows[2].cells[3], hours_str)

    # Table 3 – problem description (cols 1-3 are merged into one cell)
    desc = intervention["title"]
    if intervention["notes"]:
        desc += "\n" + intervention["notes"]
    _cell_set(doc.tables[3].rows[7].cells[1], desc)

    return doc


def _today():
    return datetime.now().strftime("%Y-%m-%d")


def _current_year_month():
    now = datetime.now()
    return now.year, now.month


def _current_week_range():
    today = date.today()
    monday = today - timedelta(days=today.weekday())
    sunday = monday + timedelta(days=6)
    return monday, sunday


# ── Dashboard ─────────────────────────────────────────────────────────────────

@app.route("/")
def dashboard():
    interventions = db.get_interventions_this_month()
    total_hours = sum(i["duration_hours"] for i in interventions)
    return render_template("dashboard.html",
                           interventions=interventions,
                           total_hours=total_hours)


# ── Clients ───────────────────────────────────────────────────────────────────

@app.route("/clients")
def clients():
    all_clients = db.get_all_clients()
    return render_template("clients.html", clients=all_clients)


@app.route("/clients/new", methods=["GET", "POST"])
def new_client():
    if request.method == "POST":
        name = request.form["name"].strip()
        notes = request.form.get("notes", "").strip()
        if name:
            db.create_client(name, notes)
        return redirect(url_for("clients"))
    return render_template("client_form.html", client=None)


@app.route("/clients/<int:client_id>/edit", methods=["GET", "POST"])
def edit_client(client_id):
    client = db.get_client(client_id)
    if client is None:
        return redirect(url_for("clients"))
    if request.method == "POST":
        name = request.form["name"].strip()
        notes = request.form.get("notes", "").strip()
        if name:
            db.update_client(client_id, name, notes)
        return redirect(url_for("clients"))
    return render_template("client_form.html", client=client)


@app.route("/clients/<int:client_id>/delete", methods=["POST"])
def delete_client(client_id):
    db.delete_client(client_id)
    return redirect(url_for("clients"))


# ── Interventions ─────────────────────────────────────────────────────────────

@app.route("/log", methods=["GET", "POST"])
def log_intervention():
    clients = db.get_all_clients()
    if request.method == "POST":
        client_id = int(request.form["client_id"])
        date = request.form["date"]
        type_ = request.form["type"]
        location = request.form["location"]
        title = request.form["title"].strip()
        duration_hours = float(request.form["duration_hours"])
        notes = request.form.get("notes", "").strip()
        db.create_intervention(client_id, date, type_, location, title, duration_hours, notes)
        return redirect(url_for("dashboard"))
    return render_template("intervention_form.html",
                           clients=clients,
                           intervention=None,
                           types=db.INTERVENTION_TYPES,
                           locations=db.LOCATIONS,
                           today=_today())


@app.route("/interventions/<int:intervention_id>/edit", methods=["GET", "POST"])
def edit_intervention(intervention_id):
    intervention = db.get_intervention(intervention_id)
    if intervention is None:
        return redirect(url_for("dashboard"))
    clients = db.get_all_clients()
    if request.method == "POST":
        client_id = int(request.form["client_id"])
        date = request.form["date"]
        type_ = request.form["type"]
        location = request.form["location"]
        title = request.form["title"].strip()
        duration_hours = float(request.form["duration_hours"])
        notes = request.form.get("notes", "").strip()
        db.update_intervention(intervention_id, client_id, date, type_, location, title, duration_hours, notes)
        return redirect(url_for("dashboard"))
    return render_template("intervention_form.html",
                           clients=clients,
                           intervention=intervention,
                           types=db.INTERVENTION_TYPES,
                           locations=db.LOCATIONS,
                           today=_today())


@app.route("/interventions/<int:intervention_id>/delete", methods=["POST"])
def delete_intervention(intervention_id):
    db.delete_intervention(intervention_id)
    return redirect(url_for("dashboard"))


# ── Week view ─────────────────────────────────────────────────────────────────

_DAY_NAMES = ["Maandag", "Dinsdag", "Woensdag", "Donderdag", "Vrijdag", "Zaterdag", "Zondag"]

@app.route("/week")
def week():
    monday, sunday = _current_week_range()
    interventions = db.get_interventions_for_week(monday.isoformat(), sunday.isoformat())

    grouped = {}
    for i in interventions:
        d = i["date"]
        if d not in grouped:
            day_obj = date.fromisoformat(d)
            grouped[d] = {
                "day_name": _DAY_NAMES[day_obj.weekday()],
                "total": 0.0,
                "items": [],
            }
        grouped[d]["total"] += i["duration_hours"]
        grouped[d]["items"].append(i)

    total_hours = sum(i["duration_hours"] for i in interventions)
    return render_template("week.html",
                           grouped=grouped,
                           monday=monday,
                           sunday=sunday,
                           total_hours=total_hours)


# ── Monthly overview ──────────────────────────────────────────────────────────

@app.route("/overview")
def overview():
    year, month = _current_year_month()
    try:
        year = int(request.args.get("year", year))
        month = int(request.args.get("month", month))
    except ValueError:
        pass

    client_id = request.args.get("client_id", "")
    try:
        client_id = int(client_id) if client_id else None
    except ValueError:
        client_id = None

    selected_type = request.args.get("type_filter", "") or None

    interventions = db.get_interventions_for_month(year, month, client_id, selected_type)
    all_clients = db.get_all_clients()

    grouped = {}
    for i in interventions:
        name = i["client_name"]
        if name not in grouped:
            grouped[name] = {"total": 0.0, "items": []}
        grouped[name]["total"] += i["duration_hours"]
        grouped[name]["items"].append(i)

    months = [
        (1, "Januari"), (2, "Februari"), (3, "Maart"), (4, "April"),
        (5, "Mei"), (6, "Juni"), (7, "Juli"), (8, "Augustus"),
        (9, "September"), (10, "Oktober"), (11, "November"), (12, "December"),
    ]

    total_hours = sum(i["duration_hours"] for i in interventions)

    return render_template("overview.html",
                           grouped=grouped,
                           year=year,
                           month=month,
                           months=months,
                           total_hours=total_hours,
                           all_clients=all_clients,
                           selected_client_id=client_id,
                           all_types=db.INTERVENTION_TYPES,
                           selected_type=selected_type)


@app.route("/overview/export")
def export_csv():
    year, month = _current_year_month()
    try:
        year = int(request.args.get("year", year))
        month = int(request.args.get("month", month))
    except ValueError:
        pass

    client_id = request.args.get("client_id", "")
    try:
        client_id = int(client_id) if client_id else None
    except ValueError:
        client_id = None

    selected_type = request.args.get("type_filter", "") or None

    interventions = db.get_interventions_for_month(year, month, client_id, selected_type)

    def _flat(text):
        """Collapse newlines so notes never break a CSV row."""
        if not text:
            return ""
        return " | ".join(line.strip() for line in text.splitlines() if line.strip())

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";", quoting=csv.QUOTE_ALL)
    writer.writerow(["Klant", "Datum", "Type", "Locatie", "Titel", "Uren", "Notities"])
    for i in interventions:
        writer.writerow([
            i["client_name"],
            i["date"],
            i["type"],
            i["location"],
            i["title"],
            str(i["duration_hours"]).replace(".", ","),
            _flat(i["notes"]),
        ])

    filename = f"interventies_{year:04d}-{month:02d}.csv"
    return Response(
        "﻿" + "sep=;\n" + output.getvalue(),  # BOM + delimiter hint for Excel
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Settings ──────────────────────────────────────────────────────────────────

@app.route("/settings")
def settings():
    return render_template("settings.html")


# ── Timesheet ─────────────────────────────────────────────────────────────────

@app.route("/timesheet")
def timesheet():
    intervention_id = request.args.get("id", type=int)
    if not intervention_id:
        return redirect(url_for("overview"))
    intervention = db.get_intervention(intervention_id)
    if not intervention:
        return redirect(url_for("overview"))
    return render_template("timesheet.html", intervention=intervention)


@app.route("/timesheet/download", methods=["POST"])
def timesheet_download():
    intervention_id = int(request.form["intervention_id"])
    intervention = db.get_intervention(intervention_id)
    if not intervention:
        return redirect(url_for("overview"))

    doc = _fill_timesheet(intervention)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_client = intervention["client_name"].replace(" ", "_")
    filename = f"Timesheet_{safe_client}_{intervention['date']}.docx"
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


# ── Startup ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    db.init_db()
    Timer(1.0, lambda: webbrowser.open("http://localhost:5000")).start()
    app.run(debug=True, use_reloader=False)
